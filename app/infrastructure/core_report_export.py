"""Word and PDF report export for the ``Керн и описание`` sheet."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory

from app.domain.lithology_attributes import LITHOLOGY_ATTRIBUTE_OPTIONS


# Adjacent labels generated on separate photo columns are one geological layer.
# A small tolerance protects against harmless floating-point rounding of depth.
MERGE_DEPTH_TOLERANCE = 0.02

def export_core_description_report(file_path: str | Path, project_title: str, wells: list[dict], format_name: str) -> None:
    """Write an editable sedimentological layer description in the selected format."""
    output = Path(file_path)
    format_name = format_name.casefold()
    if format_name == "xlsx":
        _export_xlsx(output, project_title, wells)
    elif format_name == "docx":
        _export_docx_official_form(output, project_title, wells)
    elif format_name == "pdf":
        _export_pdf_official_form(output, project_title, wells)
    else:
        raise ValueError("Поддерживаются форматы Excel (.xlsx), Word (.docx) и PDF (.pdf)")


def _split_image(image, target_ratio: float = 1.55) -> list:
    """Split very tall tablets into page-friendly pieces without downsampling."""
    if image.isNull():
        return []
    max_height = max(1, round(image.width() / max(target_ratio, 0.1)))
    if image.height() <= max_height:
        return [image]
    return [image.copy(0, top, image.width(), min(max_height, image.height() - top)) for top in range(0, image.height(), max_height)]


def _well_metadata(well: dict) -> list[tuple[str, str]]:
    settings = dict(well.get("depth_settings") or {})
    depth_range = well.get("depth_range")
    unit = str(settings.get("unit") or "м")
    if isinstance(depth_range, (tuple, list)) and len(depth_range) == 2:
        depth_text = f"{float(depth_range[0]):g} - {float(depth_range[1]):g} {unit}"
    else:
        depth_text = "не задана"
    rows = [
        ("Диапазон скважины", depth_text),
        ("Система глубин", str(well.get("depth_reference") or settings.get("coordinate_system") or "MD")),
        ("Datum", str(settings.get("datum") or "не задан")),
    ]
    if settings.get("datum_elevation") is not None:
        rows.append(("Высота datum над MSL", f"{float(settings['datum_elevation']):g} {unit}"))
    core_interval = well.get("core_interval")
    if isinstance(core_interval, (tuple, list)) and len(core_interval) == 2:
        rows.append(("Интервал керна", f"{float(core_interval[0]):g} - {float(core_interval[1]):g} {unit}"))
    return rows


def _facies_key(layer: dict) -> tuple[str, str]:
    attributes = dict(layer.get("attributes") or {})
    for field in ("Код фации", "Индекс фации"):
        value = str(attributes.get(field) or "").strip()
        if value:
            return field, value.casefold()
    return "label", str(layer.get("label") or "").strip().casefold()


def _depth_or_none(value: object) -> float | None:
    try:
        return None if value is None else float(value)
    except (TypeError, ValueError):
        return None


def _append_unique(existing: object, added: object) -> str:
    """Keep one readable value when merged columns contain the same metadata."""
    values: list[str] = []
    for source in (existing, added):
        for value in str(source or "").split(" | "):
            value = value.strip()
            if value and value not in values:
                values.append(value)
    return " | ".join(values)


def _merge_layer_attributes(target: dict, source: dict) -> None:
    for key, value in dict(source or {}).items():
        if value in (None, ""):
            continue
        if target.get(key) in (None, ""):
            target[key] = value
        elif str(target[key]) != str(value):
            target[key] = _append_unique(target[key], value)


def _merged_layers(well: dict) -> list[dict]:
    """Coalesce uninterrupted occurrences of one facies across photo columns."""
    indexed_layers = list(enumerate(well.get("layers") or []))
    indexed_layers.sort(
        key=lambda item: (
            _depth_or_none(item[1].get("depth_from")) is None,
            _depth_or_none(item[1].get("depth_from")) or 0.0,
            _depth_or_none(item[1].get("depth_to")) or 0.0,
            item[0],
        )
    )
    merged: list[dict] = []
    for _, source in indexed_layers:
        layer = dict(source)
        layer["attributes"] = dict(source.get("attributes") or {})
        start, end = _depth_or_none(layer.get("depth_from")), _depth_or_none(layer.get("depth_to"))
        if (
            merged
            and start is not None
            and end is not None
            and _facies_key(merged[-1]) == _facies_key(layer)
        ):
            previous = merged[-1]
            previous_end = _depth_or_none(previous.get("depth_to"))
            if previous_end is not None and start <= previous_end + MERGE_DEPTH_TOLERANCE:
                previous_start = _depth_or_none(previous.get("depth_from"))
                previous["depth_from"] = min(previous_start if previous_start is not None else start, start)
                previous["depth_to"] = max(previous_end, end)
                _merge_layer_attributes(previous["attributes"], layer["attributes"])
                continue
        merged.append(layer)
    return merged


def _layer_rows(well: dict) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    unit = str((well.get("depth_settings") or {}).get("unit") or "м")
    for layer in _merged_layers(well):
        start, end = layer.get("depth_from"), layer.get("depth_to")
        interval = "не задан"
        if start is not None and end is not None:
            interval = f"{float(start):g} - {float(end):g} {unit}"
        attributes = dict(layer.get("attributes") or {})
        attributes.pop("Литология", None)
        details = "; ".join(f"{key}: {value}" for key, value in attributes.items() if value not in (None, "")) or "-"
        rows.append((interval, str(layer.get("label") or "Слой"), details))
    return rows


# The field guide contains 16 lithological parameters.  The client template
# keeps them as one readable description, rather than widening the Excel form
# with sixteen technical columns.
LITHOLOGY_DESCRIPTION_FIELDS = tuple(LITHOLOGY_ATTRIBUTE_OPTIONS.keys())

OFFICIAL_HEADERS = (
    "Месторождение", "№ скв.", "№ долбления", "Интервал отбора керна, м\nКровля",
    "Интервал отбора керна, м\nПодошва", "Проходка, м", "Вынос керна, м", "Вынос, %",
    "Интервал фации, м\nКровля", "Интервал фации, м\nПодошва", "№ слоя",
    "Толщина фации, м", "Индекс фации", "Название фации", "Код фации",
    "Краткое описание", "Примечание", "Литологическое описание (16 параметров)",
)


def _text(value: object, default: str = "") -> str:
    return default if value in (None, "") else str(value)


def _lithology_description(attributes: dict[str, object]) -> str:
    """Render all selected guide parameters into one editable cell."""
    lines = [
        f"{field}: {_text(attributes.get(field))}"
        for field in LITHOLOGY_DESCRIPTION_FIELDS
        if _text(attributes.get(field))
    ]
    return "\n".join(lines)


def _official_rows(project_title: str, wells: list[dict]) -> list[list[object]]:
    """Make rows similar to the client's layered sedimentological Excel form."""
    rows: list[list[object]] = []
    for well_index, well in enumerate(wells, 1):
        settings = dict(well.get("depth_settings") or {})
        core_interval = well.get("core_interval") or well.get("depth_range")
        if not isinstance(core_interval, (tuple, list)) or len(core_interval) != 2:
            core_interval = (None, None)
        layers = _merged_layers(well)
        for layer_index, layer in enumerate(layers, 1):
            attributes = dict(layer.get("attributes") or {})
            top, base = layer.get("depth_from"), layer.get("depth_to")
            thickness = None
            if top is not None and base is not None:
                thickness = round(abs(float(base) - float(top)), 3)
            facies = _text(attributes.get("Индекс фации"), _text(layer.get("label"), "не указано"))
            name = _text(attributes.get("Название фации"), facies)
            description = _text(attributes.get("Краткое описание"), "Описание не заполнено")
            rows.append([
                _text(settings.get("field") or settings.get("месторождение"), project_title or "не указано"),
                _text(well.get("name"), str(well_index)), _text(settings.get("drilling_number")),
                core_interval[0], core_interval[1],
                round(abs(float(core_interval[1]) - float(core_interval[0])), 3)
                if core_interval[0] is not None and core_interval[1] is not None else "",
                "", "", top, base, layer_index, thickness, facies, name,
                _text(attributes.get("Код фации")), description, _text(attributes.get("Примечание")),
                _lithology_description(attributes),
            ])
    return rows


def _export_xlsx(file_path: Path, project_title: str, wells: list[dict]) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise RuntimeError("Для экспорта Excel установите openpyxl: pip install openpyxl") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Послойное описание"
    last_column = len(OFFICIAL_HEADERS)
    last_column_letter = get_column_letter(last_column)
    sheet.merge_cells(f"A1:{last_column_letter}1")
    sheet["A1"] = "ПОСЛОЙНОЕ СЕДИМЕНТОЛОГИЧЕСКОЕ ОПИСАНИЕ ПОЛНОРАЗМЕРНОГО КЕРНА СКВАЖИНЫ"
    sheet["A1"].font = Font(name="Times New Roman", size=14, bold=True)
    sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    sheet.merge_cells(f"A2:{get_column_letter(last_column - 1)}2")
    sheet["A2"] = f"Проект: {project_title or 'не указан'}"
    sheet.cell(2, last_column).value = "Приложение 2"
    sheet.cell(2, last_column).alignment = Alignment(horizontal="right")
    sheet.merge_cells("D3:E3")
    sheet["D3"] = "Интервал отбора керна, м"
    sheet.merge_cells("I3:J3")
    sheet["I3"] = "Интервал фации по бурению, м"
    for column, title in enumerate(OFFICIAL_HEADERS, 1):
        cell = sheet.cell(4, column)
        cell.value = title.split("\n")[-1] if column in {4, 5, 9, 10} else title
    for column in (column for column in range(1, last_column + 1) if column not in {4, 5, 9, 10}):
        sheet.merge_cells(start_row=3, start_column=column, end_row=4, end_column=column)
        sheet.cell(3, column).value = OFFICIAL_HEADERS[column - 1]
    header_fill = PatternFill("solid", fgColor="D9EAD3")
    thin = Side(style="thin", color="606060")
    for row in sheet.iter_rows(min_row=3, max_row=4, min_col=1, max_col=last_column):
        for cell in row:
            cell.font = Font(name="Times New Roman", size=9, bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.fill = header_fill
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    sheet.row_dimensions[1].height = 28
    sheet.row_dimensions[3].height = 30
    sheet.row_dimensions[4].height = 54
    for row_number, row in enumerate(_official_rows(project_title, wells), 5):
        for column, value in enumerate(row, 1):
            cell = sheet.cell(row_number, column, value)
            cell.font = Font(name="Times New Roman", size=9)
            cell.alignment = Alignment(horizontal="center" if column < 16 else "left", vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
        text_length = max(len(str(row[15])), len(str(row[17])))
        sheet.row_dimensions[row_number].height = max(34, min(180, 13 * (1 + text_length // 65)))
    widths = (
        18, 12, 11, 11, 11, 11, 12, 9, 11, 11, 9, 11, 14, 22, 11, 32, 24, 48,
    )
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A5"
    sheet.auto_filter.ref = f"A4:{last_column_letter}{max(5, sheet.max_row)}"
    sheet.sheet_view.showGridLines = False
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A3
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.print_title_rows = "1:4"
    sheet.page_margins.left = 0.2
    sheet.page_margins.right = 0.2
    sheet.page_margins.top = 0.4
    sheet.page_margins.bottom = 0.4
    workbook.save(file_path)


def _set_cell_shading(cell, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table, widths: list[int]) -> None:
    """Set stable, fixed Word table geometry in twips."""
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    table_pr.append(table_width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    table_pr.append(indent)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_width = cell._tc.tcPr.tcW
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _export_docx(file_path: Path, project_title: str, wells: list[dict]) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("Для экспорта Word установите python-docx: pip install python-docx") from exc

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(0.55))
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.text = "Kern Analyzer | Керн и описание"
    header.style = styles["Normal"]
    header.runs[0].font.color.rgb = RGBColor(100, 110, 130)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Сформировано Kern Analyzer")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЁТ ПО КЕРНУ И ОПИСАНИЮ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 77, 120)
    title.paragraph_format.space_after = Pt(4)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Проект: {project_title or 'Без названия'}\nСформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    subtitle.paragraph_format.space_after = Pt(14)

    with TemporaryDirectory(prefix="kern_analyzer-report-") as temp_dir:
        temp = Path(temp_dir)
        for well_index, well in enumerate(wells, 1):
            if well_index > 1:
                document.add_page_break()
            document.add_heading(f"Скважина {well.get('name') or well_index}", level=1)
            metadata = _well_metadata(well)
            table = document.add_table(rows=len(metadata), cols=2)
            _set_table_widths(table, [2600, 12700])
            for row, (label, value) in zip(table.rows, metadata):
                row.cells[0].text, row.cells[1].text = label, value
                _set_cell_shading(row.cells[0], "E8EEF5")
                for run in row.cells[0].paragraphs[0].runs:
                    run.bold = True

            fragments = _split_image(well.get("image"))
            for image_index, image in enumerate(fragments, 1):
                document.add_paragraph("Планшет «Керн и описание»" + (f", фрагмент {image_index}" if len(fragments) > 1 else ""))
                image_path = temp / f"well-{well_index}-part-{image_index}.png"
                if image.save(str(image_path), "PNG", 100):
                    document.add_picture(str(image_path), width=Inches(10.2))
                if image_index < len(fragments):
                    document.add_page_break()
                    document.add_heading(f"Скважина {well.get('name') or well_index} - продолжение", level=1)

            rows = _layer_rows(well)
            document.add_heading("Описание выделенных слоёв", level=2)
            if not rows:
                document.add_paragraph("Выделенные слои пока отсутствуют.")
            else:
                table = document.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                _set_table_widths(table, [2800, 3600, 8900])
                for cell, title_text in zip(table.rows[0].cells, ("Интервал", "Фация", "Параметры")):
                    cell.text = title_text
                    _set_cell_shading(cell, "E8EEF5")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                for layer in rows:
                    cells = table.add_row().cells
                    for cell, value in zip(cells, layer):
                        cell.text = value
    document.save(str(file_path))


def _export_pdf(file_path: Path, project_title: str, wells: list[dict]) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Image as PdfImage, KeepTogether, PageBreak, Paragraph, Spacer, Table, TableStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate
    except ImportError as exc:
        raise RuntimeError("Для экспорта PDF установите reportlab: pip install reportlab") from exc

    # DejaVu is bundled by most Python/reportlab environments and keeps Russian
    # labels intact.  If unavailable, ReportLab's standard font is still usable.
    font_name = "Helvetica"
    for candidate in (
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ):
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("Kern AnalyzerReport", str(candidate)))
            font_name = "Kern AnalyzerReport"
            break
    document = SimpleDocTemplate(
        str(file_path),
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=f"Kern Analyzer - {project_title}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Kern AnalyzerTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=22, textColor=colors.HexColor("#1f4d78"), alignment=1, spaceAfter=5 * mm)
    heading_style = ParagraphStyle("Kern AnalyzerHeading", parent=styles["Heading1"], fontName=font_name, fontSize=14, leading=17, textColor=colors.HexColor("#2e74b5"), spaceBefore=5 * mm, spaceAfter=3 * mm)
    body_style = ParagraphStyle("Kern AnalyzerBody", parent=styles["BodyText"], fontName=font_name, fontSize=8.5, leading=11, textColor=colors.HexColor("#303846"))
    small_style = ParagraphStyle("Kern AnalyzerSmall", parent=body_style, fontSize=7, leading=9)
    story = [
        Paragraph("ОТЧЁТ ПО КЕРНУ И ОПИСАНИЮ", title_style),
        Paragraph(f"Проект: {project_title or 'Без названия'}<br/>Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}", body_style),
        Spacer(1, 3 * mm),
    ]
    available_width = document.width
    with TemporaryDirectory(prefix="kern_analyzer-report-pdf-") as temp_dir:
        temp = Path(temp_dir)
        for well_index, well in enumerate(wells, 1):
            if well_index > 1:
                story.append(PageBreak())
            name = str(well.get("name") or well_index)
            story.append(Paragraph(f"Скважина {name}", heading_style))
            metadata = [[Paragraph(f"<b>{label}</b>", body_style), Paragraph(value, body_style)] for label, value in _well_metadata(well)]
            metadata_table = Table(metadata, colWidths=[40 * mm, available_width - 40 * mm], hAlign="LEFT")
            metadata_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e8eef5")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#cfd7e4")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dce3ed")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.extend([metadata_table, Spacer(1, 4 * mm)])
            fragments = _split_image(well.get("image"))
            for image_index, image in enumerate(fragments, 1):
                image_path = temp / f"well-{well_index}-part-{image_index}.png"
                if image.save(str(image_path), "PNG", 100):
                    from PIL import Image as PilImage

                    with PilImage.open(image_path) as pil:
                        ratio = min(available_width / pil.width, (150 * mm) / pil.height)
                        report_image = PdfImage(str(image_path), width=pil.width * ratio, height=pil.height * ratio)
                    caption = Paragraph("Планшет «Керн и описание»" + (f", фрагмент {image_index}" if len(fragments) > 1 else ""), small_style)
                    story.append(KeepTogether([caption, Spacer(1, 1.5 * mm), report_image, Spacer(1, 3 * mm)]))
            rows = _layer_rows(well)
            story.append(Paragraph("Описание выделенных слоёв", heading_style))
            if not rows:
                story.append(Paragraph("Выделенные слои пока отсутствуют.", body_style))
                continue
            header = [Paragraph(f"<b>{label}</b>", small_style) for label in ("Интервал", "Фация", "Параметры")]
            table_rows = [header] + [[Paragraph(value, small_style) for value in row] for row in rows]
            layers_table = Table(table_rows, colWidths=[35 * mm, 38 * mm, available_width - 73 * mm], repeatRows=1, hAlign="LEFT")
            layers_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef5")),
                ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#cfd7e4")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#dce3ed")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(layers_table)
        document.build(story)


def _export_docx_official_form(file_path: Path, project_title: str, wells: list[dict]) -> None:
    """Export the same field set as the editable Excel form to a printable Word table."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("Для экспорта Word установите python-docx: pip install python-docx") from exc

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(16.54)  # A3 landscape
    section.page_height = Inches(11.69)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(0.25))
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(7)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПОСЛОЙНОЕ СЕДИМЕНТОЛОГИЧЕСКОЕ ОПИСАНИЕ ПОЛНОРАЗМЕРНОГО КЕРНА СКВАЖИНЫ")
    run.bold = True
    run.font.size = Pt(13)
    subtitle = document.add_paragraph(f"Проект: {project_title or 'не указан'}                                                     Приложение 2")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(rows=1, cols=len(OFFICIAL_HEADERS))
    table.style = "Table Grid"
    widths = [900, 650, 580, 620, 620, 580, 620, 460, 620, 620, 440, 560, 720, 1250, 620, 2600, 1200, 3400]
    _set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, OFFICIAL_HEADERS):
        cell.text = header
        _set_cell_shading(cell, "D9EAD3")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(6)
    for row in _official_rows(project_title, wells):
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row)):
            cell.text = "" if value is None else str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index >= 15 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(6.5)
    document.save(str(file_path))


def _export_pdf_official_form(file_path: Path, project_title: str, wells: list[dict]) -> None:
    """Export a compact A3 PDF copy of the official-style sedimentological table."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A3, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.platypus import SimpleDocTemplate
    except ImportError as exc:
        raise RuntimeError("Для экспорта PDF установите reportlab: pip install reportlab") from exc

    font_name = "Helvetica"
    for candidate in (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/calibri.ttf")):
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("Kern AnalyzerOfficial", str(candidate)))
            font_name = "Kern AnalyzerOfficial"
            break
    document = SimpleDocTemplate(
        str(file_path), pagesize=landscape(A3), leftMargin=7 * mm, rightMargin=7 * mm,
        topMargin=8 * mm, bottomMargin=8 * mm, title=f"Послойное описание — {project_title}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("OfficialTitle", parent=styles["Title"], fontName=font_name, fontSize=12, leading=14, alignment=1)
    text_style = ParagraphStyle("OfficialCell", parent=styles["BodyText"], fontName=font_name, fontSize=5.4, leading=6.3, alignment=1)
    description_style = ParagraphStyle("OfficialDescription", parent=text_style, alignment=0)
    story = [
        Paragraph("ПОСЛОЙНОЕ СЕДИМЕНТОЛОГИЧЕСКОЕ ОПИСАНИЕ ПОЛНОРАЗМЕРНОГО КЕРНА СКВАЖИНЫ", title_style),
        Paragraph(f"Проект: {project_title or 'не указан'} &nbsp;&nbsp;&nbsp;&nbsp; Приложение 2", text_style),
        Spacer(1, 3 * mm),
    ]
    header = [Paragraph(f"<b>{value.replace(chr(10), '<br/>')}</b>", text_style) for value in OFFICIAL_HEADERS]
    table_data = [header]
    for row in _official_rows(project_title, wells):
        table_data.append([
            Paragraph("" if value is None else str(value).replace("\n", "<br/>"), description_style if index >= 15 else text_style)
            for index, value in enumerate(row)
        ])
    widths = [16, 12, 11, 12, 12, 11, 12, 9, 12, 12, 9, 11, 14, 24, 12, 48, 24, 65]
    table = Table(table_data, colWidths=[value * mm for value in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAD3")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#606060")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(table)
    document.build(story)
